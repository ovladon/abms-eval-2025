import streamlit as st
from Crypto.Random import get_random_bytes
from aspect_based_metadata_generator import generate_aspect_based_metadata
import json
import time
import threading
import gc
import tempfile
import os
import queue
import psutil
import pickle
import hashlib
import pandas as pd
import plotly.express as px

from PIL import Image
import wave
from moviepy.editor import VideoFileClip
import docx

from analysis_modules import (
    ActionabilityAnalysis,
    AudienceAppropriatenessAnalysis,
    CognitiveAnalysis,
    ComplexityAnalysis,
    ControversialityAnalysis,
    CulturalContextAnalysis,
    EmotionalPolarityAnalysis,
    EthicalConsiderationsAnalysis,
    FormalismAnalysis,
    GenreAnalysis,
    HumorAnalysis,
    IntentionalityAnalysis,
    InteractivityAnalysis,
    LexicalDiversityAnalysis,
    ModalityAnalysis,
    MultimodalityAnalysis,
    NarrativeStyleAnalysis,
    NoveltyAnalysis,
    ObjectivityAnalysis,
    PersuasivenessAnalysis,
    QuantitativeAnalysis,
    QualitativeAnalysis,
    ReadabilityAnalysis,
    ReliabilityAnalysis,
    SentimentAnalysis,
    SocialOrientationAnalysis,
    SpecificityAnalysis,
    SpatialAnalysis,
    SyntacticComplexityAnalysis,
    TemporalAnalysis,
)

try:
    from vosk import Model, KaldiRecognizer
except ImportError:
    st.error("Vosk library not found. Please install it using `pip install vosk`.")
    st.stop()

vosk_model = None
def load_vosk_model():
    global vosk_model
    if vosk_model is not None:
        return vosk_model

    model_path = "models/vosk-model-small-en-us-0.15"
    if not os.path.exists(model_path):
        with st.spinner("Downloading Vosk model... This may take a while."):
            os.makedirs("models", exist_ok=True)
            import urllib.request
            import zipfile
            url = "https://alphacephei.com/vosk/models/vosk-model-small-en-us-0.15.zip"
            zip_path = "models/vosk-model-small-en-us-0.15.zip"
            try:
                urllib.request.urlretrieve(url, zip_path)
                with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                    zip_ref.extractall("models")
                os.remove(zip_path)
            except Exception as e:
                st.error(f"Failed to download Vosk model: {e}")
                st.stop()
    try:
        vosk_model = Model(model_path)
        return vosk_model
    except Exception as e:
        st.error(f"Failed to load Vosk model: {e}")
        st.stop()

blip_processor = None
blip_model = None
def load_blip_model():
    global blip_processor, blip_model
    if blip_model is not None and blip_processor is not None:
        return blip_processor, blip_model
    try:
        from transformers import BlipProcessor, BlipForConditionalGeneration
        import torch
        device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
        blip_processor = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-base")
        blip_model = BlipForConditionalGeneration.from_pretrained("Salesforce/blip-image-captioning-base")
        blip_model.to(device)
        return blip_processor, blip_model
    except Exception as e:
        st.error(f"Error loading BLIP model: {e}")
        return None, None

class DataIngestion:
    def get_data(self):
        uploaded_file = st.file_uploader(
            "Upload your file",
            type=["txt", "pdf", "docx", "jpg", "jpeg", "png", "mp3", "wav", "mp4", "avi", "csv", "xlsx"]
        )
        if uploaded_file is not None:
            st.session_state.data_file_name = uploaded_file.name
            file_type = uploaded_file.type
            if "text" in file_type:
                return self.extract_text(uploaded_file), 'text'
            elif "pdf" in file_type:
                return self.extract_pdf_text(uploaded_file), 'text'
            elif "word" in file_type or uploaded_file.name.endswith('.docx'):
                return self.extract_docx_text(uploaded_file), 'text'
            elif "image" in file_type:
                return self.extract_image_description(uploaded_file), 'image'
            elif "audio" in file_type:
                return self.extract_audio_transcription(uploaded_file), 'audio'
            elif "video" in file_type:
                return self.extract_video_transcription(uploaded_file), 'video'
            elif "spreadsheet" in file_type or "csv" in file_type:
                return self.extract_structured_data_description(uploaded_file), 'structured'
            else:
                st.error("Unsupported file type.")
                return None, None
        else:
            text_input = st.text_area("Or enter text here:")
            if text_input:
                return text_input, 'text'
            else:
                return None, None

    def extract_text(self, file):
        try:
            return file.read().decode('utf-8')
        except Exception as e:
            st.error(f"Error reading text file: {e}")
            return ""

    def extract_pdf_text(self, file):
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                extracted_text = page.extract_text()
                if extracted_text:
                    text += extracted_text + " "
            return text.strip()
        except Exception as e:
            st.error(f"Error extracting text from PDF: {e}")
            return ""

    def extract_docx_text(self, file):
        try:
            import docx
            doc = docx.Document(file)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text.strip()
        except Exception as e:
            st.error(f"Error extracting text from .docx file: {e}")
            return ""

    def extract_image_description(self, file):
        try:
            image = Image.open(file)
            st.image(image, caption='Uploaded Image', use_column_width=True)
            return self.generate_image_description(image)
        except Exception as e:
            st.error(f"Error processing image: {e}")
            return ""

    def generate_image_description(self, image, processor=None, model=None):
        try:
            if processor is None or model is None:
                processor, model = load_blip_model()
                if model is None or processor is None:
                    return "An image."
            import torch
            device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
            inputs = processor(image, return_tensors="pt").to(device)
            out = model.generate(**inputs)
            description = processor.decode(out[0], skip_special_tokens=True)
            return description
        except Exception as e:
            st.error(f"Error generating image description: {e}")
            return "An image."

    def extract_audio_transcription(self, file):
        try:
            model = load_vosk_model()
            with tempfile.NamedTemporaryFile(delete=False, suffix=file.name) as tmp_file:
                tmp_file.write(file.read())
                tmp_file_path = tmp_file.name

            from pydub import AudioSegment
            audio = AudioSegment.from_file(tmp_file_path)
            audio = audio.set_channels(1)
            audio = audio.set_frame_rate(16000)

            converted_tmp_file_path = tmp_file_path + '_converted.wav'
            audio.export(converted_tmp_file_path, format='wav')

            wf = wave.open(converted_tmp_file_path, "rb")
            rec = KaldiRecognizer(model, wf.getframerate())
            rec.SetWords(True)

            transcription = ""
            while True:
                data = wf.readframes(4000)
                if len(data) == 0:
                    break
                if rec.AcceptWaveform(data):
                    result = rec.Result()
                    result_dict = json.loads(result)
                    transcription += result_dict.get("text", "") + " "
            final_result = rec.FinalResult()
            final_result_dict = json.loads(final_result)
            transcription += final_result_dict.get("text", "")
            wf.close()

            os.unlink(tmp_file_path)
            os.unlink(converted_tmp_file_path)
            return transcription.strip()
        except Exception as e:
            st.error(f"Error processing audio: {e}")
            return ""

    def extract_video_transcription(self, file):
        try:
            model = load_vosk_model()
            blip_processor, blip_model = load_blip_model()
            if blip_model is None or blip_processor is None:
                return ""

            with tempfile.NamedTemporaryFile(delete=False, suffix='.mp4') as tmp_file:
                tmp_file.write(file.read())
                tmp_file_path = tmp_file.name

            video = VideoFileClip(tmp_file_path)
            audio_path = tmp_file_path + ".wav"
            video.audio.write_audiofile(audio_path, verbose=False, logger=None)

            from pydub import AudioSegment
            audio_segment = AudioSegment.from_file(audio_path)
            audio_segment = audio_segment.set_channels(1)
            audio_segment = audio_segment.set_frame_rate(16000)
            converted_audio_path = tmp_file_path + "_converted.wav"
            audio_segment.export(converted_audio_path, format='wav')

            wf = wave.open(converted_audio_path, "rb")
            rec = KaldiRecognizer(model, wf.getframerate())
            rec.SetWords(True)

            transcription = ""
            while True:
                data = wf.readframes(4000)
                if len(data) == 0:
                    break
                if rec.AcceptWaveform(data):
                    result = rec.Result()
                    result_dict = json.loads(result)
                    transcription += result_dict.get("text", "") + " "
            final_result = rec.FinalResult()
            final_result_dict = json.loads(final_result)
            transcription += final_result_dict.get("text", "")
            wf.close()

            image_descriptions = self.extract_video_key_frame_descriptions(tmp_file_path)
            combined_text = self.combine_video_descriptions(image_descriptions, transcription)

            video.reader.close()
            video.audio.reader.close_proc()
            video.close()
            os.unlink(audio_path)
            os.unlink(converted_audio_path)
            os.unlink(tmp_file_path)
            return combined_text

        except Exception as e:
            st.error(f"Error processing video: {e}")
            return ""

    def extract_video_key_frame_descriptions(self, video_path, interval_seconds=30):
        try:
            blip_processor, blip_model = load_blip_model()
            if blip_model is None or blip_processor is None:
                return []
            video = VideoFileClip(video_path)
            duration = int(video.duration)
            image_descriptions = []
            max_frames = 5
            intervals = max(1, duration // max_frames)
            times = list(range(0, duration, intervals))
            for t in times:
                if t > duration:
                    break
                frame = video.get_frame(t)
                image = Image.fromarray(frame)
                description = self.generate_image_description(image, blip_processor, blip_model)
                image_descriptions.append(description)
            video.reader.close()
            video.audio.reader.close_proc()
            video.close()
            return image_descriptions
        except Exception as e:
            st.error(f"Error extracting key frames: {e}")
            return []

    def combine_video_descriptions(self, image_descriptions, audio_transcription):
        combined_text = " ".join(image_descriptions) + " " + audio_transcription
        return combined_text

    def extract_structured_data_description(self, file):
        try:
            if file.type == "text/csv":
                df = pd.read_csv(file)
            else:
                df = pd.read_excel(file)
            st.write("### Structured Data Preview")
            st.dataframe(df.head())
            description = f"A structured data file with {df.shape[0]} rows and {df.shape[1]} columns."
            return description
        except Exception as e:
            st.error(f"Error processing structured data: {e}")
            return ""

def main():
    st.title("Publisher Application")
    st.write("Upload your data (text, PDF, DOCX, image, audio, video, CSV, Excel) to generate the Aspect-Based Metadata.")

    if 'analysis_in_progress' not in st.session_state:
        st.session_state.analysis_in_progress = False

    data_ingestion = DataIngestion()
    text, data_type = data_ingestion.get_data()

    if text:
        if 'analysis_results' not in st.session_state:
            st.session_state.analysis_results = None
            st.session_state.aspect_based_metadata = None
            st.session_state.encryption_key = None
            st.session_state.error = None
            st.session_state.progress_file = None

        analysis_button = st.button("Start Analysis", disabled=st.session_state.analysis_in_progress)
        resume_button = st.button("Resume Analysis") if st.session_state.progress_file and not st.session_state.analysis_in_progress else None

        if analysis_button or resume_button:
            st.session_state.analysis_in_progress = True

            result_queue = queue.Queue()
            control_event = threading.Event()
            analysis_thread = threading.Thread(target=start_analysis, args=(text, data_type, result_queue, control_event))
            analysis_thread.start()

            progress_bar = st.progress(0)
            eta_placeholder = st.empty()
            status_placeholder = st.empty()
            resource_placeholder = st.empty()

            cancel_button = st.button("Cancel Analysis")
            pause_button = st.button("Pause Analysis")

            analysis_start_time = time.time()
            resource_usage_data = []

            while analysis_thread.is_alive() or not result_queue.empty():
                try:
                    message = result_queue.get_nowait()
                    if message['type'] == 'progress':
                        progress_bar.progress(message['progress'])
                        eta_placeholder.text(message['eta_text'])
                        status_placeholder.text(message['status'])
                    elif message['type'] == 'error':
                        st.error(message['content'])
                    elif message['type'] == 'complete':
                        st.success("Analysis Complete!")
                        st.session_state.analysis_results = message.get('analysis_results')
                        st.session_state.aspect_based_metadata = message.get('aspect_based_metadata')
                        st.session_state.encryption_key = message.get('encryption_key')
                    elif message['type'] == 'progress_file':
                        st.session_state.progress_file = message.get('path')
                except queue.Empty:
                    pass

                if cancel_button:
                    control_event.set()
                    control_event.action = 'cancel'
                    st.warning("Cancelling analysis...")
                    break

                if pause_button:
                    control_event.set()
                    control_event.action = 'pause'
                    st.warning("Pausing analysis...")
                    break

                current_time = time.time() - analysis_start_time
                cpu_usage = psutil.cpu_percent(interval=0.1)
                memory_usage = psutil.virtual_memory().percent
                resource_usage_data.append({"time": current_time, "cpu": cpu_usage, "memory": memory_usage})
                resource_placeholder.text(f"CPU Usage: {cpu_usage}% | Memory Usage: {memory_usage}%")
                time.sleep(0.1)

            analysis_thread.join()
            st.session_state.analysis_in_progress = False

            total_analysis_duration = time.time() - analysis_start_time
            st.session_state.total_analysis_duration = total_analysis_duration
            st.write(f"**Analysis completed in {total_analysis_duration:.2f} seconds.**")
            df_resources = pd.DataFrame(resource_usage_data)
            fig = px.line(df_resources, x="time", y=["cpu", "memory"],
                          labels={"time": "Time (s)", "value": "Usage (%)"},
                          title="System Resource Load During Analysis")
            st.session_state.resource_fig = fig
            st.plotly_chart(fig, use_container_width=True)

            st.subheader("Save Analysis Report")
            target_folder = st.text_input("Enter target folder path (e.g., Results):", value="Results")
            if st.button("Generate and Download Report"):
                if not os.path.exists(target_folder):
                    os.makedirs(target_folder)
                from docx import Document
                from docx.shared import Inches
                document = Document()
                document.add_heading("Analysis Report", 0)
                file_name = st.session_state.get("data_file_name", "unknown")
                document.add_paragraph(f"Data file name: {file_name}")
                document.add_paragraph(f"Total analysis duration: {st.session_state.total_analysis_duration:.2f} seconds")
                document.add_heading("Aspect Scores", level=1)
                aspect_limits = load_aspect_limits()
                for key, value in st.session_state.analysis_results.items():
                    if key == 'categorical_counts':
                        continue
                    aspect_name = key.replace('_', ' ').capitalize()
                    if isinstance(value, float):
                        limits = aspect_limits.get(key, (0.0, 1.0))
                        document.add_paragraph(f"{aspect_name}: {value:.2f} (Limits: {limits[0]} - {limits[1]})")
                    else:
                        document.add_paragraph(f"{aspect_name}: {value}")
                document.add_heading("Resource Usage Graph", level=1)
                image_file = os.path.join(target_folder, "resource_usage_graph.png")
                fig.write_image(image_file, format="png", scale=2)
                document.add_picture(image_file, width=Inches(6))
                report_file = os.path.join(target_folder, "analysis_report.docx")
                document.save(report_file)
                with open(report_file, "rb") as file:
                    st.download_button(
                        label="Download Analysis Report",
                        data=file,
                        file_name="analysis_report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                st.success(f"Analysis report saved to {report_file}")

            if not control_event.is_set() or control_event.action == 'resume':
                if st.session_state.error:
                    st.error(st.session_state.error)
                else:
                    if st.session_state.aspect_based_metadata and st.session_state.encryption_key:
                        st.write("**Aspect-Based Metadata:**")
                        st.code(st.session_state.aspect_based_metadata)
                        st.write("**Encryption Key (hex):**")
                        st.code(st.session_state.encryption_key.hex())
                        st.subheader("Analysis Results")
                        aspect_limits = load_aspect_limits()
                        for key, value in st.session_state.analysis_results.items():
                            if key == 'categorical_counts':
                                continue
                            aspect_name = key.replace('_', ' ').capitalize()
                            if isinstance(value, float):
                                limits = aspect_limits.get(key, (0.0, 1.0))
                                st.write(f"**{aspect_name}:** {value:.2f} (Limits: {limits[0]} - {limits[1]})")
                            else:
                                st.write(f"**{aspect_name}:** {value}")
                    else:
                        st.error("Analysis did not complete successfully. Please check for errors.")
        else:
            if st.session_state.analysis_results:
                if st.session_state.aspect_based_metadata and st.session_state.encryption_key:
                    st.write("**Aspect-Based Metadata:**")
                    st.code(st.session_state.aspect_based_metadata)
                    st.write("**Encryption Key (hex):**")
                    st.code(st.session_state.encryption_key.hex())
                    st.subheader("Analysis Results")
                    aspect_limits = load_aspect_limits()
                    for key, value in st.session_state.analysis_results.items():
                        if key == 'categorical_counts':
                            continue
                        aspect_name = key.replace('_', ' ').capitalize()
                        if isinstance(value, float):
                            limits = aspect_limits.get(key, (0.0, 1.0))
                            st.write(f"**{aspect_name}:** {value:.2f} (Limits: {limits[0]} - {limits[1]})")
                        else:
                            st.write(f"**{aspect_name}:** {value}")
                    target_folder = st.text_input("Enter target folder path (e.g., Results):", value="Results")
                    if st.button("Generate and Download Report", key="resume_save"):
                        if not os.path.exists(target_folder):
                            os.makedirs(target_folder)
                        from docx import Document
                        from docx.shared import Inches
                        document = Document()
                        document.add_heading("Analysis Report", 0)
                        file_name = st.session_state.get("data_file_name", "unknown")
                        document.add_paragraph(f"Data file name: {file_name}")
                        document.add_paragraph(f"Total analysis duration: {st.session_state.total_analysis_duration:.2f} seconds")
                        document.add_heading("Aspect Scores", level=1)
                        aspect_limits = load_aspect_limits()
                        for key, value in st.session_state.analysis_results.items():
                            if key == 'categorical_counts':
                                continue
                            aspect_name = key.replace('_', ' ').capitalize()
                            if isinstance(value, float):
                                limits = aspect_limits.get(key, (0.0, 1.0))
                                document.add_paragraph(f"{aspect_name}: {value:.2f} (Limits: {limits[0]} - {limits[1]})")
                            else:
                                document.add_paragraph(f"{aspect_name}: {value}")
                        document.add_heading("Resource Usage Graph", level=1)
                        image_file = os.path.join(target_folder, "resource_usage_graph.png")
                        st.session_state.resource_fig.write_image(image_file, format="png", scale=2)
                        document.add_picture(image_file, width=Inches(6))
                        report_file = os.path.join(target_folder, "analysis_report.docx")
                        document.save(report_file)
                        with open(report_file, "rb") as file:
                            st.download_button(
                                label="Download Analysis Report",
                                data=file,
                                file_name="analysis_report.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        st.success(f"Analysis report saved to {report_file}")
                else:
                    st.error("Analysis did not complete successfully. Please check for errors.")
    else:
        st.write("Please upload a file or enter text.")

def start_analysis(text, data_type, result_queue, control_event):
    try:
        analysis_modules = get_analysis_modules(data_type)
        number_of_modules = len(analysis_modules)
        chunk_size = 1000
        num_chunks = len(text) // chunk_size + 1
        total_steps = number_of_modules * num_chunks

        if 'progress_file' in st.session_state and st.session_state.progress_file:
            progress_file_path = st.session_state.progress_file
            with open(progress_file_path, 'rb') as f:
                analysis_results, categorical_counts, start_step = pickle.load(f)
        else:
            analysis_results = {}
            categorical_counts = {}
            start_step = 0
            progress_file_path = None

        step = start_step
        start_time = time.time()
        base_sleep_time = 0.1
        current_sleep_time = base_sleep_time
        max_attempts = 5

        def cpu_dummy_work(n=100000):
            x = 0
            for i in range(n):
                x += i
            return x

        for chunk_index in range(num_chunks):
            if control_event.is_set():
                break
            chunk = text[chunk_index * chunk_size : (chunk_index + 1) * chunk_size]
            if not chunk:
                continue
            for module in analysis_modules:
                if control_event.is_set():
                    break
                try:
                    analysis_instance = module(chunk)
                    result = analysis_instance.analyze()
                    analysis_results = aggregate_results(analysis_results, result, categorical_counts)
                except Exception as e:
                    result_queue.put({'type': 'error', 'content': f"Error in module {module.__name__}: {e}"})
                    control_event.set()
                    control_event.action = 'error'
                    break
                step += 1
                progress = step / total_steps
                elapsed_time = time.time() - start_time
                eta = (elapsed_time / step) * (total_steps - step) if step > 0 else 0
                eta_text = f"Estimated time remaining: {int(eta)} seconds"
                status_text = f"Processing module {module.__name__} (Step {step}/{total_steps})"
                result_queue.put({
                    'type': 'progress',
                    'progress': progress,
                    'eta_text': eta_text,
                    'status': status_text
                })
                if step % 10 == 0:
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.pkl') as tmp_file:
                        pickle.dump((analysis_results, categorical_counts, step), tmp_file)
                        progress_file_path = tmp_file.name
                        result_queue.put({'type': 'progress_file', 'path': progress_file_path})
                gc.collect()
                for attempt in range(max_attempts):
                    cpu_usage = psutil.cpu_percent(interval=0.1)
                    memory_usage = psutil.virtual_memory().percent
                    if cpu_usage > 70 or memory_usage > 70:
                        current_sleep_time *= 2
                        if attempt == max_attempts - 1 and chunk_size > 100:
                            chunk_size //= 2
                            num_chunks = len(text) // chunk_size + 1
                            total_steps = number_of_modules * num_chunks
                        time.sleep(current_sleep_time)
                        cpu_dummy_work(20000)
                    else:
                        current_sleep_time = base_sleep_time
                        break
            del chunk
            gc.collect()

        if not control_event.is_set():
            data_hash = hashlib.sha256(text.encode('utf-8')).hexdigest()
            analysis_results["data_hash"] = data_hash
            encryption_key = get_random_bytes(16)
            aspect_based_metadata = generate_aspect_based_metadata(analysis_results, encryption_key)
            if aspect_based_metadata is None:
                result_queue.put({'type': 'error', 'content': "Error generating the Aspect-Based Metadata."})
            else:
                result_queue.put({
                    'type': 'complete',
                    'analysis_results': analysis_results,
                    'aspect_based_metadata': aspect_based_metadata,
                    'encryption_key': encryption_key
                })
            if progress_file_path and os.path.exists(progress_file_path):
                os.remove(progress_file_path)
                result_queue.put({'type': 'progress_file', 'path': None})
        elif control_event.action == 'pause':
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pkl') as tmp_file:
                pickle.dump((analysis_results, categorical_counts, step), tmp_file)
                progress_file_path = tmp_file.name
                result_queue.put({'type': 'progress_file', 'path': progress_file_path})
    except Exception as e:
        result_queue.put({'type': 'error', 'content': f"An error occurred during analysis: {e}"})

def get_analysis_modules(data_type):
    return [
        ActionabilityAnalysis,
        AudienceAppropriatenessAnalysis,
        CognitiveAnalysis,
        ComplexityAnalysis,
        ControversialityAnalysis,
        CulturalContextAnalysis,
        EmotionalPolarityAnalysis,
        EthicalConsiderationsAnalysis,
        FormalismAnalysis,
        GenreAnalysis,
        HumorAnalysis,
        IntentionalityAnalysis,
        InteractivityAnalysis,
        LexicalDiversityAnalysis,
        ModalityAnalysis,
        MultimodalityAnalysis,
        NarrativeStyleAnalysis,
        NoveltyAnalysis,
        ObjectivityAnalysis,
        PersuasivenessAnalysis,
        QuantitativeAnalysis,
        QualitativeAnalysis,
        ReadabilityAnalysis,
        ReliabilityAnalysis,
        SentimentAnalysis,
        SocialOrientationAnalysis,
        SpecificityAnalysis,
        SpatialAnalysis,
        SyntacticComplexityAnalysis,
        TemporalAnalysis,
    ]

def aggregate_results(existing_results, new_results, categorical_counts):
    for key, value in new_results.items():
        if key in existing_results:
            if isinstance(value, float):
                existing_results[key] = (existing_results[key] + value) / 2
            elif isinstance(value, str):
                if key not in categorical_counts:
                    categorical_counts[key] = {}
                if value in categorical_counts[key]:
                    categorical_counts[key][value] += 1
                else:
                    categorical_counts[key][value] = 1
                most_frequent = max(categorical_counts[key], key=categorical_counts[key].get)
                existing_results[key] = most_frequent
        else:
            existing_results[key] = value
    return existing_results

def load_aspect_limits():
    return {
        'actionability_analysis': (0.0, 1.0),
        'audience_appropriateness_analysis': (0.0, 3.0),
        'cognitive_analysis': (0.0, 20.0),
        'complexity_analysis': (0.0, 1.0),
        'controversiality_analysis': (0.0, 1.0),
        'cultural_context_analysis': (0.0, 1.0),
        'emotional_polarity_analysis': (0.0, 1.0),
        'ethical_considerations_analysis': (0.0, 2.0),
        'formalism_analysis': (0.0, 1.0),
        'genre_analysis': (0.0, 9.0),
        'humor_analysis': (0.0, 1.0),
        'intentionality_analysis': (0.0, 5.0),
        'interactivity_analysis': (0.0, 1.0),
        'lexical_diversity_analysis': (0.0, 1.0),
        'modality_analysis': (0.0, 3.0),
        'multimodality_analysis': (0.0, 4.0),
        'narrative_style_analysis': (0.0, 2.0),
        'novelty_analysis': (0.0, 1.0),
        'objectivity_analysis': (0.0, 1.0),
        'persuasiveness_analysis': (0.0, 1.0),
        'quantitative_analysis': (0.0, 1.0),
        'qualitative_analysis': (0.0, 1.0),
        'readability_analysis': (0.0, 100.0),
        'reliability_analysis': (0.0, 1.0),
        'sentiment_analysis': (-1.0, 1.0),
        'social_orientation_analysis': (0.0, 1.0),
        'specificity_analysis': (0.0, 1.0),
        'syntactic_complexity_analysis': (0.0, 1.0),
        'temporal_analysis': (0.0, 1.0),
    }

if __name__ == "__main__":
    main()

