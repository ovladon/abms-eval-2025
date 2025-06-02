# reader/result_presenter.py

from content_speculation import ContentSpeculator
import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import plotly.io as pio
import os

class ResultPresenter:
    def __init__(self, analysis_results):
        self.analysis_results = analysis_results
        self.aspect_descriptions = self.load_aspect_descriptions()
        self.aspect_limits = self.load_aspect_limits()

    def load_aspect_descriptions(self):
        # Descriptions for each aspect with method explanations.
        return {
            'actionability_analysis': 'Measures how action-oriented the text is.\n\n**Method:** Counts imperative sentences.',
            'audience_appropriateness_analysis': 'Determines the suitable audience level.\n\n**Method:** Uses readability scores.',
            'cognitive_analysis': 'Assesses cognitive load.\n\n**Method:** Multiple readability indices.',
            'complexity_analysis': 'Structural complexity.\n\n**Method:** Sentence length and syntax analysis.',
            'controversiality_analysis': 'Identifies controversy.\n\n**Method:** Topic modeling.',
            'cultural_context_analysis': 'Detects cultural references.\n\n**Method:** Named Entity Recognition for cultural entities.',
            'emotional_polarity_analysis': 'Determines the emotional tone.\n\n**Method:** Sentiment analysis models.',
            'ethical_considerations_analysis': 'Evaluates ethical implications.\n\n**Method:** Checks against ethical guidelines.',
            'formalism_analysis': 'Assesses the formality of language.\n\n**Method:** Formality analysis based on vocabulary and style.',
            'genre_analysis': 'Classifies the text into a genre.\n\n**Method:** Genre classification models.',
            'humor_analysis': 'Detects humorous content.\n\n**Method:** Analysis of humor-related linguistic features.',
            'intentionality_analysis': 'Identifies the primary intent of the text.\n\n**Method:** Intent recognition models.',
            'interactivity_analysis': 'Measures interactivity in the text.\n\n**Method:** Counts direct addresses and questions.',
            'lexical_diversity_analysis': 'Evaluates vocabulary variety.\n\n**Method:** Type-token ratio and other metrics.',
            'modality_analysis': 'Identifies the sensory modality of content.\n\n**Method:** Detects textual, visual, or auditory cues.',
            'multimodality_analysis': 'Checks for multiple content types.\n\n**Method:** Looks for mentions of various media types.',
            'narrative_style_analysis': 'Assesses narrative perspective.\n\n**Method:** Identifies first-, second-, or third-person usage.',
            'novelty_analysis': 'Measures the originality of the text.\n\n**Method:** Compares against known texts for uniqueness.',
            'objectivity_analysis': 'Assesses objectivity.\n\n**Method:** Detects subjective language markers.',
            'persuasiveness_analysis': 'Evaluates persuasiveness.\n\n**Method:** Looks for persuasive language patterns.',
            'quantitative_analysis': 'Extracts numerical data.\n\n**Method:** Statistical extraction and analysis.',
            'qualitative_analysis': 'Assesses descriptive richness.\n\n**Method:** Analysis of descriptive language.',
            'readability_analysis': 'Measures readability.\n\n**Method:** Readability scoring algorithms.',
            'reliability_analysis': 'Checks reliability of information.\n\n**Method:** Verification against authoritative sources.',
            'sentiment_analysis': 'Analyzes sentiment.\n\n**Method:** Sentiment analysis models.',
            'social_orientation_analysis': 'Assesses focus on individual versus collective language.\n\n**Method:** Pronoun usage analysis.',
            'specificity_analysis': 'Measures detail and specificity.\n\n**Method:** Analysis of descriptive precision.',
            'spatial_analysis': 'Detects spatial references.\n\n**Method:** Named Entity Recognition for locations.',
            'syntactic_complexity_analysis': 'Evaluates syntactic complexity.\n\n**Method:** Analysis of sentence structures and parse trees.',
            'temporal_analysis': 'Identifies temporal references.\n\n**Method:** Verb tense and time expression analysis.'
        }

    def load_aspect_limits(self):
        return {
            'actionability_analysis': (0.0, 1.0),
            'cognitive_analysis': (0.0, 20.0),
            'complexity_analysis': (0.0, 1.0),
            'controversiality_analysis': (0.0, 1.0),
            'emotional_polarity_analysis': (0.0, 1.0),
            'formalism_analysis': (0.0, 1.0),
            'humor_analysis': (0.0, 1.0),
            'interactivity_analysis': (0.0, 1.0),
            'lexical_diversity_analysis': (0.0, 1.0),
            'novelty_analysis': (0.0, 1.0),
            'objectivity_analysis': (0.0, 1.0),
            'persuasiveness_analysis': (0.0, 1.0),
            'quantitative_analysis': (0.0, 1.0),
            'qualitative_analysis': (0.0, 1.0),
            'readability_analysis': (0.0, 100.0),
            'reliability_analysis': (0.0, 1.0),
            'sentiment_analysis': (-1.0, 1.0),
            'social_orientation_analysis': (0.0, 1.0),
            'specificity_analysis': (0.0, 1.0),
            'syntactic_complexity_analysis': (0.0, 1.0),
            'temporal_analysis': (0.0, 1.0),
        }

    def present_in_streamlit(self):
        st.header("Analysis Results")

        # Explicitly display the data hash if it exists.
        if "data_hash" in self.analysis_results:
            st.write("**Data Hash:**", self.analysis_results["data_hash"])

        # Separate numerical and categorical results
        numerical_results = {k: v for k, v in self.analysis_results.items() if isinstance(v, float)}
        categorical_results = {k: v for k, v in self.analysis_results.items() if not isinstance(v, float)}

        st.subheader("Analysis Overview")
        combined_results = []
        for key in self.analysis_results:
            # Avoid re-displaying the hash in the table since it's already shown above.
            if key == "data_hash":
                continue
            aspect_name = key.replace('_', ' ').capitalize()
            value = self.analysis_results[key]
            description = self.aspect_descriptions.get(key, '')
            if isinstance(value, float):
                limits = self.aspect_limits.get(key, (0.0, 1.0))
                result_str = f"{value:.2f} (Limits: {limits[0]} - {limits[1]})"
            else:
                result_str = str(value)
            combined_results.append({
                'Aspect': aspect_name,
                'Result': result_str,
                'Description': description
            })

        df_combined = pd.DataFrame(combined_results).astype(str)
        st.dataframe(df_combined)

        # Visualization
        self.visualize_results(numerical_results, categorical_results)

        if st.button("Generate Report"):
            self.generate_report()

    def visualize_results(self, numerical_results, categorical_results):
        st.subheader("Visualization of Analysis Results")

        category_mapping = self.get_category_mappings()
        all_results = numerical_results.copy()
        for k, v in categorical_results.items():
            mapping = category_mapping.get(k)
            if mapping:
                numerical_value = mapping.get(v, 0)
                all_results[k] = numerical_value

        aspect_order = [
            'actionability_analysis',
            'audience_appropriateness_analysis',
            'cognitive_analysis',
            'complexity_analysis',
            'controversiality_analysis',
            'cultural_context_analysis',
            'emotional_polarity_analysis',
            'ethical_considerations_analysis',
            'formalism_analysis',
            'genre_analysis',
            'humor_analysis',
            'intentionality_analysis',
            'interactivity_analysis',
            'lexical_diversity_analysis',
            'modality_analysis',
            'multimodality_analysis',
            'narrative_style_analysis',
            'novelty_analysis',
            'objectivity_analysis',
            'persuasiveness_analysis',
            'quantitative_analysis',
            'qualitative_analysis',
            'readability_analysis',
            'reliability_analysis',
            'sentiment_analysis',
            'social_orientation_analysis',
            'specificity_analysis',
            'spatial_analysis',
            'syntactic_complexity_analysis',
            'temporal_analysis'
        ]

        df_all = pd.DataFrame({
            'Aspect': [key.replace('_', ' ').capitalize() for key in aspect_order],
            'Score': [all_results.get(key, 0) for key in aspect_order]
        })

        df_all = df_all.iloc[::-1]

        fig = go.Figure(go.Bar(
            x=df_all['Score'],
            y=df_all['Aspect'],
            orientation='h',
            text=df_all['Score'],
            textposition='auto'
        ))
        fig.update_layout(
            margin=dict(l=300, r=50, t=50, b=50),
            width=1600,
            height=1200
        )
        st.plotly_chart(fig, use_container_width=True)
        st.markdown("**Legend for Categorical Values:**")
        category_mapping_display = self.get_category_mappings_display()
        for key, mapping in category_mapping_display.items():
            st.write(f"**{key}:**")
            legend_items = ', '.join([f"{k} = {v}" for k, v in mapping.items()])
            st.write(legend_items)

        self.fig_for_report = fig

    def get_category_mappings(self):
        return {
            'audience_appropriateness_analysis': {'Children': 0, 'Middle School': 1, 'High School': 2, 'Adult': 3},
            'cultural_context_analysis': {'General': 0, 'Cultural Specific': 1},
            'ethical_considerations_analysis': {'Low': 0, 'Medium': 1, 'High': 2},
            'genre_analysis': {'Political Speech': 0, 'News': 1, 'Story': 2, 'Academic': 3, 'Legal': 4, 'Scientific': 5, 'Finance': 6, 'Entertainment': 7, 'Sports': 8, 'Historical Document': 9},
            'intentionality_analysis': {'Informative': 0, 'Persuasive': 1, 'Narrative': 2, 'Descriptive': 3, 'Expository': 4, 'Instructional': 5},
            'modality_analysis': {'Textual': 0, 'Visual': 1, 'Auditory': 2, 'Multimedia': 3},
            'multimodality_analysis': {'text': 0, 'image': 1, 'audio': 2, 'video': 3, 'interactive': 4},
            'narrative_style_analysis': {'First_Person': 0, 'Second_Person': 1, 'Third_Person': 2},
            'spatial_analysis': {'General': 0, 'Local': 1, 'Regional': 2, 'Global': 3}
        }

    def get_category_mappings_display(self):
        mappings = self.get_category_mappings()
        display_mappings = {}
        for key, mapping in mappings.items():
            display_key = key.replace('_', ' ').capitalize()
            inverse_map = {v: k for k, v in mapping.items()}
            display_mappings[display_key] = inverse_map
        return display_mappings

    def generate_report(self):
        results_folder = 'Results'
        os.makedirs(results_folder, exist_ok=True)

        from docx import Document
        from docx.shared import Inches

        document = Document()
        document.add_heading('Analysis Report', 0)

        document.add_heading('Analysis Results', level=1)
        # Display the data hash at the beginning of the report if present
        if "data_hash" in self.analysis_results:
            document.add_paragraph(f"Data Hash: {self.analysis_results['data_hash']}")

        for key, value in self.analysis_results.items():
            if key == "data_hash":
                continue  # Already displayed above
            aspect_name = key.replace('_', ' ').capitalize()
            if isinstance(value, float):
                limits = self.aspect_limits.get(key, (0.0, 1.0))
                document.add_paragraph(f"{aspect_name}: {value:.2f} (Limits: {limits[0]} - {limits[1]})")
            else:
                document.add_paragraph(f"{aspect_name}: {value}")

        document.add_heading('Legend for Categorical Values', level=1)
        category_mapping_display = self.get_category_mappings_display()
        for key, mapping in category_mapping_display.items():
            document.add_heading(key, level=2)
            for num_value, category in mapping.items():
                document.add_paragraph(f"{num_value}: {category}", style='List Bullet')

        from content_speculation import ContentSpeculator
        speculator = ContentSpeculator(self.analysis_results)

        document.add_heading('Content Speculation', level=1)
        matched_categories = speculator.speculate_content(output_format='human')
        if matched_categories:
            for category in matched_categories:
                document.add_heading(f"Inferred Content Category: {category['name']}", level=2)
                document.add_paragraph(f"Description: {category['human_description']}")
                document.add_paragraph(f"Related Keywords: {', '.join(category['keywords'])}")
                document.add_paragraph(f"Confidence Score: {category['confidence_score']}")
                document.add_paragraph("References:")
                for ref in category.get('references', []):
                    document.add_paragraph(f"- {ref['reference']}: {ref['reference_explanation']}", style='List Bullet')
                document.add_paragraph()
        else:
            document.add_paragraph("No inferred content categories could be determined from the analysis results.")

        document.add_heading('Aspect Synergies', level=1)
        synergy_results = speculator.speculate_synergies(output_format='human')
        if synergy_results:
            for synergy in synergy_results:
                document.add_heading(f"Inferred Synergy: {synergy['name']}", level=2)
                document.add_paragraph(f"Description: {synergy['human_description']}")
                document.add_paragraph("Aspects Involved: " + ', '.join(synergy['aspects_involved']))
                document.add_paragraph(f"Confidence Score: {synergy['confidence_score']}")
                document.add_paragraph("References:")
                for ref in synergy.get('references', []):
                    document.add_paragraph(f"- {ref['reference']}: {ref['reference_explanation']}", style='List Bullet')
                document.add_paragraph()
        else:
            document.add_paragraph("No significant inferred synergies could be determined from the current analysis.")

        image_file = os.path.join(results_folder, 'analysis_results.png')
        self.fig_for_report.update_layout(width=1600, height=1200, margin=dict(l=300))
        self.fig_for_report.write_image(image_file, format='png', scale=3)
        document.add_heading('Visualization of Analysis Results', level=1)
        document.add_picture(image_file, width=Inches(9))
        report_file = os.path.join(results_folder, 'analysis_report.docx')
        document.save(report_file)
        st.success(f"Report saved to {report_file}")

# End of file

